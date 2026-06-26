"""文章路由 — /api/articles/*, /api/author-articles/*, /api/linked-articles/*, /api/records/*/reanalyze*

含三大板块：
  - 公众号文章（articles）：同步、列表、添加、详情、下载图片、分析估值
  - 作者文章（author-articles）：导入、列表、爬取全文、URL 提取
  - 个人文档（linked-articles）：上传、下载、内容预览、embedding、分块
  - 分析记录（records）：重新分析单张图片
"""

import asyncio
import json
import logging
import os
import re
from pathlib import Path

logger = logging.getLogger(__name__)

from fastapi import APIRouter, HTTPException, UploadFile, BackgroundTasks
from fastapi.responses import FileResponse

from config import ROOT, IMAGES_DIR, UPLOADS_DIR
from state import (
    analyze_progress as _analyze_progress,
    analyze_cancel as _analyze_cancel,
    analyze_tasks as _analyze_tasks,
    reanalyze_tasks as _reanalyze_tasks,
    vision_semaphore as _vision_semaphore,
    crawl_semaphore as _crawl_semaphore,
)
from db import (
    sync_articles, list_articles, get_article, get_article_by_url, create_article,
    update_article, create_analysis_record, update_analysis_record,
    get_analysis_records, get_analysis_record,
    list_all_analysis_records,
    create_author_article, update_author_article, get_author_article_by_url,
    list_author_articles, get_author_article, delete_author_article, count_author_articles,
    create_linked_article, list_linked_articles, get_linked_article, delete_linked_article,
    update_linked_article_file, update_linked_article_embed_status,
    save_document_chunks, get_document_chunks,
    save_valuation,
)
from article_reader import fetch_article, download_images
from image_parser import ImageParser
from rag import index_to_chroma, index_author_article
from models.articles import ExtractUrlRequest
from routers.tasks import CreateTaskRequest

router = APIRouter(tags=["articles"])


# ══════════════════════════════════════════════════════
# 公众号文章 API
# ══════════════════════════════════════════════════════

@router.post("/api/articles/sync")
async def sync_articles_api():
    """从 articles.json 同步文章列表到 DB。"""
    articles_file = ROOT / "doc" / "articles.json"
    if not articles_file.exists():
        raise HTTPException(400, "articles.json 不存在")
    with open(articles_file) as f:
        articles = json.load(f)
    sync_articles(articles)
    count = len(list_articles())
    return {"ok": True, "total": count}


@router.get("/api/articles")
async def list_articles_api(status: str = None):
    """文章列表，可选按状态筛选。"""
    return {"articles": list_articles(status)}


@router.post("/api/articles/add")
async def add_article_api(req: CreateTaskRequest, background_tasks: BackgroundTasks):
    """粘贴公众号链接，自动解析+下载+分析（全流程）。"""
    url = req.url.strip()
    if not url:
        raise HTTPException(400, "链接不能为空")

    # 去重检查
    existing = get_article_by_url(url)
    if existing:
        # 如果之前的状态是 pending 或 error，允许重新触发
        if existing["status"] in ("pending", "error"):
            article_id = existing["id"]
            update_article(article_id, status="pending", title="解析中...", error_msg=None)
            background_tasks.add_task(_background_add_article, article_id, url)
            return {"ok": True, "message": "已重新提交解析", "article_id": article_id}
        return {"ok": False, "message": "文章已存在", "article_id": existing["id"], "status": existing["status"]}

    # 先插入占位记录
    article_id = create_article(url, title="解析中...")

    # 后台全流程：解析 → 下载 → 分析（使用 FastAPI BackgroundTasks 确保任务执行）
    background_tasks.add_task(_background_add_article, article_id, url)
    return {"ok": True, "message": "已提交，正在解析", "article_id": article_id}


@router.get("/api/articles/{article_id}")
async def get_article_api(article_id: int):
    """文章详情，含图片列表和分析记录。"""
    article = get_article(article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    records = get_analysis_records(article_id)
    article["analysis_records"] = records
    return article


@router.post("/api/articles/{article_id}/download")
async def download_article_images(article_id: int):
    """异步下载文章中的图片（后台执行）。"""
    article = get_article(article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    if article["status"] == "downloading":
        return {"ok": False, "message": "正在下载中，请稍候"}

    update_article(article_id, status="downloading")
    asyncio.create_task(_background_download(article_id, article))
    return {"ok": True, "message": "下载已开始"}


@router.post("/api/articles/{article_id}/analyze")
async def analyze_article_images(article_id: int):
    """异步分析文章中的所有图片（后台执行）。"""
    article = get_article(article_id)
    if not article:
        raise HTTPException(404, "文章不存在")

    records = get_analysis_records(article_id)
    pending = [r for r in records if r["status"] in ("pending", "error", "cancelled", "timeout")]
    if not pending:
        return {"ok": True, "message": "没有待分析的图片", "total": 0}
    if article["status"] == "analyzing":
        return {"ok": False, "message": "正在分析中，请稍候"}

    update_article(article_id, status="analyzing")
    _analyze_progress[article_id] = {"total": len(pending), "done": 0, "success": 0, "failed": 0, "current_record_id": None}
    asyncio.create_task(_background_analyze(article_id))
    return {"ok": True, "message": "分析已开始"}


@router.get("/api/articles/{article_id}/analyze-status")
async def get_analyze_status(article_id: int):
    """查询后台分析任务进度。"""
    article = get_article(article_id)
    if not article:
        return {"status": "unknown", "progress": {}}

    # 防卡死检测：文章显示 analyzing 但无活跃后台任务且所有记录已完成
    if article["status"] == "analyzing" and article_id not in _analyze_progress:
        records = get_analysis_records(article_id)
        pending = [r for r in records if r["status"] in ("pending", "analyzing", "error", "cancelled", "timeout")]
        if not pending:
            total_success = len([r for r in records if r["status"] == "success"])
            if total_success > 0:
                update_article(article_id, status="analyzed", error_msg="")
                article["status"] = "analyzed"
            else:
                update_article(article_id, status="error", error_msg="所有图片分析失败")
                article["status"] = "error"

    progress = _analyze_progress.get(article_id, {})
    return {
        "status": article["status"],
        "progress": progress,
    }


@router.post("/api/articles/{article_id}/cancel-analyze")
async def cancel_analyze(article_id: int):
    """取消正在进行的分析任务。直接操作数据库，不依赖后台任务。"""
    article = get_article(article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    if article["status"] != "analyzing":
        return {"ok": False, "message": "当前没有在分析中"}

    # 1. 设置取消标志 + 取消 asyncio Task
    _analyze_cancel.add(article_id)
    task = _analyze_tasks.get(article_id)
    if task and not task.done():
        task.cancel()

    # 2. 直接标记所有 pending/error 记录为 cancelled
    records = get_analysis_records(article_id)
    for r in records:
        if r["status"] in ("pending", "error"):
            update_analysis_record(r["id"], status="cancelled", error_msg="用户取消分析")

    # 3. 更新文章状态为 error
    update_article(article_id, status="error", error_msg="用户取消分析")

    # 4. 清理进度跟踪
    _analyze_progress.pop(article_id, None)

    return {"ok": True, "message": "已取消"}


# ── 分析记录重新分析 API ─────────────────────────────────

@router.post("/api/records/{record_id}/reanalyze")
async def reanalyze_image(record_id: int):
    """重新分析单张图片（异步，后台执行，前端轮询查状态）。"""
    record = get_analysis_record(record_id)
    if not record:
        raise HTTPException(404, "记录不存在")

    # 已有分析正在进行则直接返回
    task = _reanalyze_tasks.get(record_id)
    if task and not task.done():
        return {"ok": True, "message": "分析已在进行中"}

    # 重置记录状态为 analyzing
    update_analysis_record(record_id, status="analyzing", error_msg="")

    # 后台启动
    bt = asyncio.create_task(_background_reanalyze(record_id))
    _reanalyze_tasks[record_id] = bt
    return {"ok": True, "message": "分析已开始", "record_id": record_id}


@router.get("/api/records/{record_id}/reanalyze-status")
async def get_reanalyze_status(record_id: int):
    """查询单张图片重新分析的状态。"""
    record = get_analysis_record(record_id)
    if not record:
        raise HTTPException(404, "记录不存在")
    task = _reanalyze_tasks.get(record_id)
    return {
        "record_id": record_id,
        "status": record["status"],
        "error_msg": record.get("error_msg"),
        "running": bool(task and not task.done()),
    }


# ══════════════════════════════════════════════════════
# 作者文章 API
# ══════════════════════════════════════════════════════

@router.post("/api/author-articles/import")
async def import_author_articles():
    """从 Excel 导入作者文章（幂等，跳过已存在）。"""
    from import_articles import import_from_excel
    result = import_from_excel()
    return {"ok": True, **result}


@router.get("/api/author-articles")
async def list_author_articles_api(status: str = None, search: str = None, limit: int = 200):
    """作者文章列表。"""
    articles = list_author_articles(status=status, search=search, limit=limit)
    stats = count_author_articles()
    return {"articles": articles, "stats": stats}


@router.get("/api/author-articles/{article_id}")
async def get_author_article_api(article_id: int):
    """作者文章详情。"""
    article = get_author_article(article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    return article


@router.delete("/api/author-articles/{article_id}")
async def delete_author_article_api(article_id: int):
    """删除作者文章。"""
    if not delete_author_article(article_id):
        raise HTTPException(404, "文章不存在")
    return {"ok": True}


@router.post("/api/author-articles/extract")
async def extract_article_from_url(body: ExtractUrlRequest):
    """从 URL 提取文章信息（通用网页，不依赖 Playwright）。"""
    import requests as req
    from bs4 import BeautifulSoup

    try:
        resp = req.get(body.url, headers={
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
        }, timeout=15)
        resp.raise_for_status()
        resp.encoding = resp.apparent_encoding or "utf-8"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求失败: {e}")

    soup = BeautifulSoup(resp.text, "lxml")

    # Extract title
    title = ""
    og_title = soup.find("meta", property="og:title")
    if og_title and og_title.get("content"):
        title = og_title["content"]
    elif soup.title and soup.title.string:
        title = soup.title.string.strip()
    elif soup.find("h1"):
        title = soup.find("h1").get_text(strip=True)

    # Extract author
    author = ""
    meta_author = soup.find("meta", attrs={"name": "author"})
    if meta_author and meta_author.get("content"):
        author = meta_author["content"]
    og_site = soup.find("meta", property="og:site_name")
    if not author and og_site and og_site.get("content"):
        author = og_site["content"]

    # Extract publish time
    publish_time = ""
    for attr in [("property", "article:published_time"), ("name", "pubdate"),
                 ("name", "publish_time"), ("property", "og:article:published_time")]:
        tag = soup.find("meta", attrs=dict([attr]))
        if tag and tag.get("content"):
            publish_time = tag["content"][:19]
            break
    if not publish_time:
        time_tag = soup.find("time")
        if time_tag:
            publish_time = time_tag.get("datetime", "")[:19] or time_tag.get_text(strip=True)[:19]

    # Extract summary
    summary = ""
    meta_desc = soup.find("meta", attrs={"name": "description"})
    if meta_desc and meta_desc.get("content"):
        summary = meta_desc["content"]
    elif soup.find("meta", property="og:description"):
        summary = soup.find("meta", property="og:description").get("content", "")

    # Extract content text (main body)
    content_text = ""
    for selector in ["article", '[role="main"]', "main", ".article-content",
                     ".post-content", ".entry-content", "#content", ".content"]:
        container = soup.select_one(selector)
        if container and len(container.get_text(strip=True)) > 100:
            content_text = container.get_text("\n", strip=True)
            break
    if not content_text:
        paragraphs = soup.find_all("p")
        content_text = "\n\n".join(p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 20)

    if len(content_text) > 5000:
        content_text = content_text[:5000] + "..."

    return {
        "url": body.url,
        "title": title,
        "author": author,
        "publish_time": publish_time,
        "summary": summary[:500] if summary else "",
        "content_text": content_text,
    }


@router.post("/api/author-articles")
async def create_author_article_api(body: dict):
    """直接创建作者文章记录。"""
    article_id = create_author_article(
        url=body.get("url", ""),
        title=body.get("title", ""),
        publish_time=body.get("publish_time", ""),
        summary=body.get("summary", ""),
        article_type=body.get("article_type", ""),
    )
    if body.get("content_text"):
        update_author_article(article_id, content_text=body["content_text"])
    return {"id": article_id}


@router.post("/api/author-articles/crawl")
async def crawl_all_author_articles():
    """批量爬取所有 pending 状态的作者文章全文。"""
    pending = list_author_articles(status="pending", limit=500)
    if not pending:
        return {"ok": True, "message": "没有待爬取的文章", "total": 0}

    asyncio.create_task(_batch_crawl_author_articles(pending))
    return {"ok": True, "message": f"开始爬取 {len(pending)} 篇文章", "total": len(pending)}


@router.post("/api/author-articles/{article_id}/crawl")
async def crawl_single_author_article(article_id: int):
    """爬取单篇作者文章全文。"""
    article = get_author_article(article_id)
    if not article:
        raise HTTPException(404, "文章不存在")

    update_author_article(article_id, status="crawling", error_msg="")
    asyncio.create_task(_crawl_one_author_article(article_id, article["url"]))
    return {"ok": True, "message": "开始爬取"}


# ══════════════════════════════════════════════════════
# 个人文档 API
# ══════════════════════════════════════════════════════

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".doc"}


@router.get("/api/linked-articles")
async def list_linked_articles_api(limit: int = 200):
    return list_linked_articles(limit=limit)


@router.post("/api/linked-articles")
async def upload_document(file: UploadFile):
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {ext}，仅支持 .txt / .md / .pdf / .docx")

    content = await file.read()
    file_size = len(content)
    title = Path(file.filename).stem

    article_id = create_linked_article(
        title=title, file_path="", file_size=file_size, file_type=ext.lstrip("."),
    )

    safe_name = f"{article_id}_{file.filename}"
    save_path = UPLOADS_DIR / safe_name
    save_path.write_bytes(content)

    update_linked_article_file(article_id, safe_name)

    # 异步 embedding（不阻塞响应）
    asyncio.create_task(_embed_linked_doc(article_id, ext.lstrip("."), content, title))

    return {"id": article_id}


@router.get("/api/linked-articles/{article_id}/download")
async def download_document(article_id: int):
    article = get_linked_article(article_id)
    if not article or not article.get("file_path"):
        raise HTTPException(status_code=404, detail="文档不存在")
    file_path = UPLOADS_DIR / article["file_path"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件已被删除")
    return FileResponse(
        str(file_path),
        filename=article["title"] + "." + article.get("file_type", ""),
        media_type="application/octet-stream",
    )


@router.get("/api/linked-articles/{article_id}/content")
async def get_document_content(article_id: int):
    article = get_linked_article(article_id)
    if not article or not article.get("file_path"):
        raise HTTPException(status_code=404, detail="文档不存在")
    file_path = UPLOADS_DIR / article["file_path"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件已被删除")

    file_type = article.get("file_type", "")
    content = ""

    try:
        if file_type in ("txt", "md"):
            content = file_path.read_text(encoding="utf-8", errors="replace")
        elif file_type == "pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(str(file_path))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            content = "\n\n".join(pages)
        elif file_type == "docx":
            from docx import Document
            doc = Document(str(file_path))
            content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif file_type == "doc":
            content = "（.doc 格式暂不支持在线预览，请下载后查看）"
        else:
            content = "（不支持的文件格式）"
    except Exception as e:
        content = f"内容提取失败: {e}"

    return {"content": content, "file_type": file_type}


@router.delete("/api/linked-articles/{article_id}")
async def delete_linked_article_api(article_id: int):
    article = get_linked_article(article_id)
    if not article:
        raise HTTPException(status_code=404, detail="文档不存在")
    # Delete physical file
    if article.get("file_path"):
        file_path = UPLOADS_DIR / article["file_path"]
        if file_path.exists():
            file_path.unlink()
    delete_linked_article(article_id)
    return {"ok": True}


@router.post("/api/linked-articles/{article_id}/embed")
async def embed_document(article_id: int):
    """对单篇文档做 embedding 并存入向量库。"""
    article = get_linked_article(article_id)
    if not article or not article.get("file_path"):
        raise HTTPException(status_code=404, detail="文档不存在")

    file_path = UPLOADS_DIR / article["file_path"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件已被删除")

    # 标记为 embedding 中
    update_linked_article_embed_status(article_id, "embedding")

    file_type = article.get("file_type", "")
    content = ""

    try:
        if file_type in ("txt", "md"):
            content = file_path.read_text(encoding="utf-8", errors="replace")
        elif file_type == "pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(str(file_path))
            pages = [p.extract_text() for p in reader.pages if p.extract_text()]
            content = "\n\n".join(pages)
        elif file_type == "docx":
            from docx import Document
            doc = Document(str(file_path))
            content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif file_type == "doc":
            update_linked_article_embed_status(article_id, "failed")
            raise HTTPException(status_code=400, detail=".doc 格式暂不支持 embedding，请先转换为 .docx")
    except HTTPException:
        raise
    except Exception as e:
        update_linked_article_embed_status(article_id, "failed")
        raise HTTPException(status_code=500, detail=f"内容提取失败: {e}")

    if not content.strip():
        update_linked_article_embed_status(article_id, "failed")
        raise HTTPException(status_code=400, detail="文档内容为空，无法索引")

    try:
        # 分块并存入 ChromaDB
        chunks_count = index_to_chroma("linked_doc", str(article_id), article.get("title", ""), content)

        # 保存分块到 SQLite（用于展示）
        from rag import _chunk_text
        chunks = _chunk_text(content)
        save_document_chunks(article_id, chunks)

        # 更新状态
        update_linked_article_embed_status(article_id, "done", chunks_count)

        return {"ok": True, "chunks_indexed": chunks_count}
    except Exception as e:
        update_linked_article_embed_status(article_id, "failed")
        raise HTTPException(status_code=500, detail=f"Embedding 失败: {e}")


@router.get("/api/linked-articles/{article_id}/chunks")
async def get_document_chunks_api(article_id: int):
    """获取文档的分块详情。"""
    article = get_linked_article(article_id)
    if not article:
        raise HTTPException(status_code=404, detail="文档不存在")
    chunks = get_document_chunks(article_id)
    return {"chunks": chunks, "total": len(chunks)}


# ══════════════════════════════════════════════════════
# 后台任务
# ══════════════════════════════════════════════════════

def _clean_article_html(html: str) -> str:
    """清理微信文章 HTML，保留图文混排结构。"""
    # data-src → src（微信图片用 data-src 做懒加载）
    html = html.replace('data-src="', 'src="')
    # 移除 script/style 标签
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
    # 移除 noscript
    html = re.sub(r'<noscript[^>]*>.*?</noscript>', '', html, flags=re.DOTALL | re.IGNORECASE)
    # 移除微信特有的 data- 属性
    html = re.sub(r'\s+data-[a-z-]+="[^"]*"', '', html, flags=re.IGNORECASE)
    # 移除 onclick 等事件
    html = re.sub(r'\s+on[a-z]+="[^"]*"', '', html, flags=re.IGNORECASE)
    # 图片 src 走代理（微信图片）
    html = re.sub(
        r'src="(https?://mmbiz\.qpic\.cn/[^"]*)"',
        lambda m: f'src="/api/proxy-image?url={m.group(1)}"',
        html
    )
    # 移除空的 section/div（减少嵌套）
    html = re.sub(r'<(section|div)[^>]*>\s*</\1>', '', html, flags=re.IGNORECASE)
    return html.strip()


async def _background_add_article(article_id: int, url: str):
    """全流程后台任务：解析文章 → 下载图片 → 分析估值。"""
    logger.info(f"[article:{article_id}] 开始后台解析: {url[:50]}...")
    try:
        # 1. 解析文章
        update_article(article_id, status="downloading")
        result = await fetch_article(url)
        title = result.get("title", "未知标题")
        publish_time = result.get("publish_time", "")
        images = result.get("images", [])

        if not images:
            update_article(article_id, status="error", error_msg="文章无图片", title=title, publish_time=publish_time)
            return

        # 更新文章信息
        update_article(article_id, title=title, publish_time=publish_time)

        # 2. 下载图片
        from datetime import datetime
        article_date = publish_time[:10] if publish_time else datetime.now().strftime("%Y-%m-%d")
        safe_title = "".join(c for c in title[:20] if c.isalnum() or c in " _-").strip() or "article"
        save_dir = IMAGES_DIR / article_date / safe_title
        save_dir.mkdir(parents=True, exist_ok=True)

        local_paths = await download_images(images, str(save_dir))

        manifest = {
            "url": url, "title": title, "article_date": article_date,
            "images": [{"index": i, "url": u, "local_path": str(lp)} for i, (u, lp) in enumerate(zip(images, local_paths))],
        }
        manifest_path = save_dir / "manifest.json"
        with open(manifest_path, "w") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)

        for i, (u, lp) in enumerate(zip(images, local_paths)):
            create_analysis_record(article_id, i, str(lp), u)

        images_dir_str = str(save_dir)
        try:
            images_dir_str = str(save_dir.relative_to(ROOT))
        except ValueError:
            pass

        update_article(article_id,
            status="downloaded", images_dir=images_dir_str,
            manifest_path=str(manifest_path), image_count=len(local_paths),
        )

        # 3. 自动分析
        update_article(article_id, status="analyzing")
        records = get_analysis_records(article_id)
        pending = [r for r in records if r["status"] in ("pending", "error", "cancelled", "timeout")]
        parser = ImageParser()
        success = 0
        failed = 0

        for record in pending:
            rid = record["id"]
            img_path = record["image_path"]
            if not os.path.isabs(img_path):
                img_path = str(ROOT / img_path)
            if not os.path.exists(img_path):
                update_analysis_record(rid, status="error", error_msg="文件不存在")
                failed += 1
                continue

            try:
                async with _vision_semaphore:
                    parse_result = await asyncio.wait_for(
                        asyncio.to_thread(parser.parse, img_path), timeout=600,
                    )
                has_value = parse_result and parse_result.get("index_code") and parse_result.get("current_value") is not None
                if has_value:
                    save_valuation(parse_result, source_image=img_path, source_url=url, snapshot_date=article_date)
                    update_analysis_record(rid, status="success",
                        index_code=parse_result.get("index_code"), index_name=parse_result.get("index_name"),
                        metric_type=parse_result.get("metric_type"),
                        raw_response=json.dumps(parse_result, ensure_ascii=False))
                    success += 1
                else:
                    update_analysis_record(rid, status="error",
                        index_code=parse_result.get("index_code") if parse_result else None,
                        error_msg="AI 未能提取到数据",
                        raw_response=json.dumps(parse_result, ensure_ascii=False) if parse_result else "")
                    failed += 1
            except asyncio.TimeoutError:
                update_analysis_record(rid, status="timeout", error_msg="分析超时")
                failed += 1
            except Exception as e:
                update_analysis_record(rid, status="error", error_msg=str(e))
                failed += 1

        err_msg = f"{failed} 张失败" if failed > 0 else None
        update_article(article_id, status="analyzed", error_msg=err_msg)
        logger.info(f"[article:{article_id}] 解析完成: success={success}, failed={failed}")

    except Exception as e:
        logger.error(f"[article:{article_id}] 后台任务异常: {e}")
        update_article(article_id, status="error", error_msg=str(e))


async def _background_download(article_id: int, article: dict):
    """后台下载任务。"""
    try:
        result = await fetch_article(article["url"])
        images = result["images"]
        if not images:
            update_article(article_id, status="error", error_msg="无图片")
            return

        from datetime import datetime
        article_date = article.get("publish_time", "")[:10] or datetime.now().strftime("%Y-%m-%d")
        safe_title = "".join(c for c in (article["title"] or "")[:20] if c.isalnum() or c in " _-").strip() or "article"
        save_dir = IMAGES_DIR / article_date / safe_title
        save_dir.mkdir(parents=True, exist_ok=True)

        local_paths = await download_images(images, str(save_dir))

        manifest = {
            "url": article["url"],
            "title": article["title"],
            "article_date": article_date,
            "images": [
                {"index": i, "url": url, "local_path": str(lp)}
                for i, (url, lp) in enumerate(zip(images, local_paths))
            ],
        }
        manifest_path = save_dir / "manifest.json"
        with open(manifest_path, "w") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)

        for i, (url, lp) in enumerate(zip(images, local_paths)):
            create_analysis_record(article_id, i, str(lp), url)

        # 归一化 images_dir 为相对路径（相对于项目根）
        images_dir_str = str(save_dir)
        try:
            images_dir_str = str(save_dir.relative_to(ROOT))
        except ValueError:
            pass  # 不在 ROOT 下时保留原路径

        update_article(article_id,
            status="downloaded",
            images_dir=images_dir_str,
            manifest_path=str(manifest_path),
            image_count=len(local_paths),
        )
    except Exception as e:
        update_article(article_id, status="error", error_msg=str(e))


async def _background_analyze(article_id: int):
    """后台分析任务。单张图片超时 10 分钟，取消时立即中断当前图片。"""
    IMAGE_TIMEOUT = 600  # 10 分钟
    current_task = None  # 当前图片的 asyncio Task，用于取消中断

    def _cancel_remaining(records_list, from_index):
        """标记从 from_index 开始的所有待处理记录为已取消。"""
        for r in records_list[from_index:]:
            if r["status"] in ("pending", "error"):
                update_analysis_record(r["id"], status="cancelled", error_msg="用户取消分析")

    try:
        article = get_article(article_id)
        records = get_analysis_records(article_id)
        pending = [r for r in records if r["status"] in ("pending", "error", "cancelled", "timeout")]

        parser = ImageParser()
        success = 0
        failed = 0

        for idx, record in enumerate(pending):
            # 循环开头检查取消
            if article_id in _analyze_cancel:
                _analyze_cancel.discard(article_id)
                _cancel_remaining(pending, idx)
                update_article(article_id, status="error", error_msg="用户取消分析")
                break

            rid = record["id"]
            img_path = record["image_path"]
            # 相对路径转绝对路径（相对于项目根目录）
            if not os.path.isabs(img_path):
                img_path = str(ROOT / img_path)

            # 记录当前正在分析的图片
            _analyze_progress[article_id]["current_record_id"] = rid

            if not os.path.exists(img_path):
                update_analysis_record(rid, status="error", error_msg="文件不存在")
                failed += 1
                _analyze_progress[article_id]["done"] += 1
                _analyze_progress[article_id]["failed"] = failed
                continue

            try:
                # 用 create_task 包装，这样可以被 task.cancel() 立即中断（信号量限制并发）
                async def _parse_with_semaphore():
                    async with _vision_semaphore:
                        return await asyncio.to_thread(parser.parse, img_path)

                current_task = asyncio.create_task(_parse_with_semaphore())
                _analyze_tasks[article_id] = current_task
                result = await asyncio.wait_for(current_task, timeout=IMAGE_TIMEOUT)

                has_value = (
                    result
                    and result.get("index_code")
                    and result.get("current_value") is not None
                )
                if has_value:
                    vid = save_valuation(
                        result,
                        source_image=img_path,
                        source_url=article["url"],
                        snapshot_date=article.get("publish_time", "")[:10] or None,
                    )
                    update_analysis_record(rid,
                        status="success",
                        index_code=result.get("index_code"),
                        index_name=result.get("index_name"),
                        metric_type=result.get("metric_type"),
                        raw_response=json.dumps(result, ensure_ascii=False),
                    )
                    success += 1
                else:
                    st = "success" if result.get("index_code") else "error"
                    update_analysis_record(rid,
                        status=st,
                        index_code=result.get("index_code"),
                        index_name=result.get("index_name"),
                        metric_type=result.get("metric_type"),
                        error_msg="AI 返回空值（无当前值）" if result.get("index_code") else "AI 未能提取到数据",
                        raw_response=json.dumps(result, ensure_ascii=False) if result else "",
                    )
                    failed += 1
            except asyncio.TimeoutError:
                update_analysis_record(rid, status="timeout", error_msg=f"分析超时（{IMAGE_TIMEOUT // 60}分钟）")
                failed += 1
            except asyncio.CancelledError:
                # 被用户取消
                update_analysis_record(rid, status="cancelled", error_msg="用户取消分析")
                _analyze_cancel.discard(article_id)
                _cancel_remaining(pending, idx + 1)
                update_article(article_id, status="error", error_msg="用户取消分析")
                break
            except Exception as e:
                update_analysis_record(rid, status="error", error_msg=str(e))
                failed += 1
            finally:
                current_task = None
                _analyze_tasks.pop(article_id, None)

            _analyze_progress[article_id]["done"] += 1
            _analyze_progress[article_id]["success"] = success
            _analyze_progress[article_id]["failed"] = failed
            _analyze_progress[article_id]["current_record_id"] = None

            # 每处理完一张后检查取消
            if article_id in _analyze_cancel:
                _analyze_cancel.discard(article_id)
                _cancel_remaining(pending, idx + 1)
                update_article(article_id, status="error", error_msg="用户取消分析")
                break

            await asyncio.sleep(0.3)

        # 仅当未被取消时才根据结果设置文章状态
        if article_id not in _analyze_cancel:
            all_records = get_analysis_records(article_id)
            total_success = len([r for r in all_records if r["status"] == "success"])
            if total_success > 0:
                update_article(article_id, status="analyzed", error_msg="")
            elif failed > 0:
                update_article(article_id, status="error", error_msg=f"{failed} 张图片分析失败")
    except Exception as e:
        update_article(article_id, status="error", error_msg=str(e))
    finally:
        _analyze_progress.pop(article_id, None)
        _analyze_tasks.pop(article_id, None)


async def _background_reanalyze(record_id: int):
    """后台重新分析单张图片（10分钟超时，不阻塞 event loop）。"""
    IMAGE_TIMEOUT = 600
    try:
        record = get_analysis_record(record_id)
        if not record:
            update_analysis_record(record_id, status="error", error_msg="记录不存在")
            return

        img_path = record["image_path"]
        if not os.path.isabs(img_path):
            img_path = str(ROOT / img_path)
        if not os.path.exists(img_path):
            update_analysis_record(record_id, status="error", error_msg="图片文件不存在")
            return

        article = get_article(record["article_id"])
        parser = ImageParser()

        # 异步执行 parser.parse，不阻塞 event loop（信号量限制并发）
        async with _vision_semaphore:
            result = await asyncio.wait_for(
                asyncio.to_thread(parser.parse, img_path),
                timeout=IMAGE_TIMEOUT,
            )

        has_value = (
            result
            and result.get("index_code")
            and result.get("current_value") is not None
        )
        if not has_value:
            err = "AI 返回空值（无当前值）" if result and result.get("index_code") else "AI 未能提取到数据"
            update_analysis_record(record_id,
                status="error",
                error_msg=err,
                raw_response=json.dumps(result, ensure_ascii=False) if result else "",
            )
            return

        vid = save_valuation(
            result,
            source_image=img_path,
            source_url=article["url"] if article else None,
            snapshot_date=article.get("publish_time", "")[:10] if article else None,
        )
        update_analysis_record(record_id,
            status="success",
            index_code=result.get("index_code"),
            index_name=result.get("index_name"),
            metric_type=result.get("metric_type"),
            raw_response=json.dumps(result, ensure_ascii=False),
        )
    except asyncio.TimeoutError:
        update_analysis_record(record_id, status="timeout", error_msg="分析超时（10分钟）")
    except Exception as e:
        update_analysis_record(record_id, status="error", error_msg=str(e))
    finally:
        _reanalyze_tasks.pop(record_id, None)


async def _batch_crawl_author_articles(articles: list):
    """后台批量爬取作者文章。"""
    for a in articles:
        update_author_article(a["id"], status="crawling")

    async def _crawl_with_limit(article_id, url):
        async with _crawl_semaphore:
            await _crawl_one_author_article(article_id, url)

    tasks = [_crawl_with_limit(a["id"], a["url"]) for a in articles]
    await asyncio.gather(*tasks, return_exceptions=True)


async def _crawl_one_author_article(article_id: int, url: str):
    """爬取单篇作者文章并更新数据库。"""
    try:
        result = await fetch_article(url)
        content = result.get("content_text", "")
        raw_html = result.get("content_html", "")
        title = result.get("title", "")
        publish_time = result.get("publish_time", "")
        images = result.get("images", [])

        # 清理 HTML：data-src → src，过滤脚本/样式/空白标签
        clean_html = _clean_article_html(raw_html) if raw_html else ""

        update_author_article(article_id,
            content_text=content,
            content_html=clean_html,
            title=title or None,
            publish_time=publish_time or None,
            images=json.dumps(images, ensure_ascii=False) if images else None,
            status="done",
            error_msg="",
        )

        # 索引到 RAG
        if content:
            index_author_article(article_id, title, content, publish_time or "")

    except Exception as e:
        update_author_article(article_id, status="error", error_msg=str(e)[:500])


async def _embed_linked_doc(article_id: int, file_type: str, raw_content: bytes, title: str):
    """后台任务：提取文档文本并 embedding。"""
    update_linked_article_embed_status(article_id, "embedding")
    try:
        text = ""
        if file_type in ("txt", "md"):
            text = raw_content.decode("utf-8", errors="replace")
        elif file_type == "pdf":
            from PyPDF2 import PdfReader
            import io
            reader = PdfReader(io.BytesIO(raw_content))
            pages = [p.extract_text() for p in reader.pages if p.extract_text()]
            text = "\n\n".join(pages)
        elif file_type == "docx":
            from docx import Document
            import io
            doc = Document(io.BytesIO(raw_content))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text.strip():
            chunks_count = index_to_chroma("linked_doc", str(article_id), title, text)
            from rag import _chunk_text
            chunks = _chunk_text(text)
            save_document_chunks(article_id, chunks)
            update_linked_article_embed_status(article_id, "done", chunks_count)
        else:
            update_linked_article_embed_status(article_id, "failed")
    except Exception:
        update_linked_article_embed_status(article_id, "failed")
